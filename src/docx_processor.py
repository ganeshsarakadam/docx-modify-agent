"""
Core DOCX processing module for editing documents while preserving formatting
"""

import io
import re
from typing import Dict, List, Any, Optional, Union, BinaryIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .exceptions import (
    DocxProcessingError,
    InvalidDocxFileError,
    EditOperationError,
    TextNotFoundError,
    StylePreservationError
)


class DocxProcessor:
    """
    Main class for processing DOCX files while preserving formatting and styles
    """
    
    def __init__(self):
        self.document: Optional[Document] = None
        self.original_styles: Dict[str, Any] = {}
    
    def load_from_buffer(self, file_buffer: BinaryIO) -> None:
        """
        Load a DOCX document from a file buffer
        
        Args:
            file_buffer: Binary file buffer containing DOCX data
            
        Raises:
            InvalidDocxFileError: If the file is not a valid DOCX document
        """
        try:
            # Reset the buffer position to the beginning
            file_buffer.seek(0)
            self.document = Document(file_buffer)
            self._extract_styles()
        except Exception as e:
            raise InvalidDocxFileError(f"Failed to load DOCX document: {str(e)}")
    
    def _extract_styles(self) -> None:
        """Extract and store document styles for preservation"""
        if not self.document:
            return
            
        self.original_styles = {}
        
        # Extract paragraph styles
        for paragraph in self.document.paragraphs:
            if paragraph.style:
                style_name = paragraph.style.name
                if style_name not in self.original_styles:
                    self.original_styles[style_name] = {
                        'type': 'paragraph',
                        'style': paragraph.style
                    }
    
    def find_and_replace_text(
        self, 
        search_text: str, 
        replace_text: str, 
        preserve_formatting: bool = True,
        case_sensitive: bool = True
    ) -> int:
        """
        Find and replace text in the document while preserving formatting
        
        Args:
            search_text: Text to search for
            replace_text: Text to replace with
            preserve_formatting: Whether to preserve the original formatting
            case_sensitive: Whether the search should be case sensitive
            
        Returns:
            Number of replacements made
            
        Raises:
            EditOperationError: If the replacement operation fails
        """
        if not self.document:
            raise EditOperationError("No document loaded")
        
        try:
            replacements_made = 0
            
            # Process paragraphs
            for paragraph in self.document.paragraphs:
                replacements_made += self._replace_text_in_paragraph(
                    paragraph, search_text, replace_text, preserve_formatting, case_sensitive
                )
            
            # Process tables
            for table in self.document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replacements_made += self._replace_text_in_paragraph(
                                paragraph, search_text, replace_text, preserve_formatting, case_sensitive
                            )
            
            return replacements_made
            
        except Exception as e:
            raise EditOperationError(f"Failed to replace text: {str(e)}")
    
    def _replace_text_in_paragraph(
        self, 
        paragraph: Paragraph, 
        search_text: str, 
        replace_text: str,
        preserve_formatting: bool,
        case_sensitive: bool
    ) -> int:
        """
        Replace text within a single paragraph while preserving formatting
        
        Args:
            paragraph: The paragraph to process
            search_text: Text to search for
            replace_text: Text to replace with
            preserve_formatting: Whether to preserve formatting
            case_sensitive: Whether the search should be case sensitive
            
        Returns:
            Number of replacements made in this paragraph
        """
        if not paragraph.runs:
            return 0
        
        # Get the full text of the paragraph
        full_text = paragraph.text
        
        # Perform case-insensitive search if needed
        if case_sensitive:
            search_pattern = re.escape(search_text)
        else:
            search_pattern = re.escape(search_text)
            full_text_lower = full_text.lower()
            search_text_lower = search_text.lower()
        
        replacements_made = 0
        
        if case_sensitive:
            if search_text not in full_text:
                return 0
        else:
            if search_text_lower not in full_text_lower:
                return 0
        
        if preserve_formatting:
            # Complex replacement preserving formatting across runs
            replacements_made = self._replace_with_formatting_preservation(
                paragraph, search_text, replace_text, case_sensitive
            )
        else:
            # Simple replacement
            if case_sensitive:
                new_text = full_text.replace(search_text, replace_text)
            else:
                new_text = re.sub(re.escape(search_text), replace_text, full_text, flags=re.IGNORECASE)
            
            if new_text != full_text:
                # Clear existing runs and add new text with default formatting
                paragraph.clear()
                paragraph.add_run(new_text)
                replacements_made = full_text.count(search_text) if case_sensitive else len(re.findall(re.escape(search_text), full_text, re.IGNORECASE))
        
        return replacements_made
    
    def _replace_with_formatting_preservation(
        self, 
        paragraph: Paragraph, 
        search_text: str, 
        replace_text: str,
        case_sensitive: bool
    ) -> int:
        """
        Replace text while preserving formatting across runs
        
        This is a complex operation that maintains the original formatting
        of each character in the search text when replacing it.
        """
        replacements_made = 0
        
        # Build a map of character positions to runs and their formatting
        char_to_run_map = []
        char_position = 0
        
        for run in paragraph.runs:
            run_text = run.text
            for char in run_text:
                char_to_run_map.append({
                    'char': char,
                    'run': run,
                    'position': char_position,
                    'formatting': self._extract_run_formatting(run)
                })
                char_position += 1
        
        # Find all occurrences of search_text
        full_text = ''.join([item['char'] for item in char_to_run_map])
        
        if case_sensitive:
            pattern = re.escape(search_text)
            flags = 0
        else:
            pattern = re.escape(search_text)
            flags = re.IGNORECASE
        
        matches = list(re.finditer(pattern, full_text, flags))
        
        if not matches:
            return 0
        
        # Process matches in reverse order to maintain positions
        for match in reversed(matches):
            start_pos = match.start()
            end_pos = match.end()
            
            # Get the formatting from the first character of the match
            first_char_formatting = char_to_run_map[start_pos]['formatting']
            
            # Replace the characters
            new_chars = []
            for i, char in enumerate(replace_text):
                new_chars.append({
                    'char': char,
                    'formatting': first_char_formatting,
                    'position': start_pos + i
                })
            
            # Remove old characters and insert new ones
            char_to_run_map[start_pos:end_pos] = new_chars
            replacements_made += 1
        
        # Rebuild the paragraph with preserved formatting
        if replacements_made > 0:
            self._rebuild_paragraph_with_formatting(paragraph, char_to_run_map)
        
        return replacements_made
    
    def _extract_run_formatting(self, run: Run) -> Dict[str, Any]:
        """Extract formatting information from a run"""
        return {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color.rgb else None,
        }
    
    def _apply_run_formatting(self, run: Run, formatting: Dict[str, Any]) -> None:
        """Apply formatting to a run"""
        if formatting['bold'] is not None:
            run.bold = formatting['bold']
        if formatting['italic'] is not None:
            run.italic = formatting['italic']
        if formatting['underline'] is not None:
            run.underline = formatting['underline']
        if formatting['font_name']:
            run.font.name = formatting['font_name']
        if formatting['font_size']:
            run.font.size = formatting['font_size']
        if formatting['font_color']:
            run.font.color.rgb = formatting['font_color']
    
    def _apply_bullet_formatting(self, paragraph: Paragraph) -> None:
        """Apply proper indentation and spacing for bullet points"""
        from docx.shared import Inches, Pt
        
        # Use List Paragraph style for consistent formatting (matches reference resume)
        try:
            paragraph.style = paragraph.part.document.styles['List Paragraph']
        except KeyError:
            # Fallback if List Paragraph style doesn't exist
            pass
        
        # Set left indentation for bullet points (0.5 inch from left margin)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        
        # Set hanging indent to align bullet text properly (-0.25 inch for bullet symbol)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        
        # Add some space after each bullet point for better readability
        paragraph.paragraph_format.space_after = Pt(3)
        
        # Ensure consistent line spacing
        paragraph.paragraph_format.line_spacing = 1.15
    
    def _rebuild_paragraph_with_formatting(
        self, 
        paragraph: Paragraph, 
        char_to_run_map: List[Dict[str, Any]]
    ) -> None:
        """Rebuild paragraph with preserved formatting"""
        # Clear existing runs
        paragraph.clear()
        
        if not char_to_run_map:
            return
        
        # Group consecutive characters with the same formatting
        current_text = ""
        current_formatting = char_to_run_map[0]['formatting']
        
        for item in char_to_run_map:
            if item['formatting'] == current_formatting:
                current_text += item['char']
            else:
                # Add run with current formatting
                if current_text:
                    new_run = paragraph.add_run(current_text)
                    self._apply_run_formatting(new_run, current_formatting)
                
                # Start new run
                current_text = item['char']
                current_formatting = item['formatting']
        
        # Add the last run
        if current_text:
            new_run = paragraph.add_run(current_text)
            self._apply_run_formatting(new_run, current_formatting)

    def replace_with_bullet_points(
        self, 
        search_text: str, 
        bullet_items: List[str], 
        preserve_formatting: bool = True
    ) -> int:
        """
        Replace a placeholder with properly formatted bullet points
        This method looks for template bullet points to copy their formatting
        
        Args:
            search_text: Text to search for (e.g., "{{highlights1}}")
            bullet_items: List of strings to convert to bullet points
            preserve_formatting: Whether to preserve the original formatting
            
        Returns:
            Number of replacements made
        """
        if not self.document or not bullet_items:
            return 0
        
        replacements_made = 0
        template_bullet_style = self._find_template_bullet_formatting()
        
        # Find paragraphs containing the search text
        for paragraph in self.document.paragraphs:
            if search_text in paragraph.text:
                # Find the position of the placeholder in the paragraph
                full_text = paragraph.text
                placeholder_start = full_text.find(search_text)
                
                if placeholder_start != -1:
                    # Instead of clearing the entire paragraph, use the formatting-preserving replacement
                    count = self._replace_with_formatting_preservation(
                        paragraph, search_text, f"• {bullet_items[0]}", False
                    )
                    
                    if count > 0:
                        replacements_made += count
                        
                        # Apply proper indentation and spacing to the first bullet point
                        self._apply_bullet_formatting(paragraph)
                        
                        # Add additional bullet points as new paragraphs after this one
                        parent_element = paragraph._element.getparent()
                        paragraph_index = list(parent_element).index(paragraph._element)
                        
                        for i, item in enumerate(bullet_items[1:], 1):
                            # Create new paragraph for each additional bullet point
                            new_paragraph = self.document.add_paragraph()
                            bullet_run = new_paragraph.add_run(f"• {item}")
                            if template_bullet_style:
                                self._apply_run_formatting(bullet_run, template_bullet_style)
                            
                            # Apply proper indentation and spacing for bullet points
                            self._apply_bullet_formatting(new_paragraph)
                            
                            # Insert the new paragraph after the current one
                            parent_element.insert(paragraph_index + i, new_paragraph._element)
                    
                    break  # Only replace the first occurrence
        
        return replacements_made

    def _find_template_bullet_formatting(self) -> Optional[Dict[str, Any]]:
        """
        Find a template bullet point to extract its formatting
        Looks for bullet points that contain placeholder text like "Add highlights with this style"
        """
        for paragraph in self.document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if (paragraph_text.startswith("•") and 
                "add highlights with this style" in paragraph_text.lower()):
                # Found template bullet point, extract formatting from its first run
                if paragraph.runs:
                    return self._extract_run_formatting(paragraph.runs[0])
        
        return None

    def remove_placeholder_bullet_points(self, placeholder_text: str = "Add highlights with this style.") -> int:
        """
        Remove placeholder bullet points that were used for formatting reference
        This removes lines like: "•	   Add highlights with this style."
        
        Args:
            placeholder_text: The text part of the placeholder bullet point to search for
            
        Returns:
            Number of placeholder bullet points removed
        """
        if not self.document:
            return 0
        
        removals_made = 0
        paragraphs_to_remove = []
        
        for paragraph in self.document.paragraphs:
            paragraph_text = paragraph.text.strip()
            
            # Check if this paragraph contains the placeholder bullet point
            if (paragraph_text.startswith("•") and 
                placeholder_text.lower() in paragraph_text.lower()):
                paragraphs_to_remove.append(paragraph)
                removals_made += 1
        
        # Remove the identified paragraphs
        for paragraph in paragraphs_to_remove:
            # Get the paragraph's parent element
            p_element = paragraph._element
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
        
        return removals_made

    def insert_text_at_bookmark(self, bookmark_name: str, text: str) -> bool:
        """
        Insert text at a specific bookmark location
        
        Args:
            bookmark_name: Name of the bookmark
            text: Text to insert
            
        Returns:
            True if bookmark was found and text inserted, False otherwise
        """
        if not self.document:
            raise EditOperationError("No document loaded")
        
        # This is a simplified implementation
        # For full bookmark support, you'd need to work with the XML structure
        return False
    
    def add_paragraph(self, text: str, style: Optional[str] = None) -> None:
        """
        Add a new paragraph to the document
        
        Args:
            text: Text content for the paragraph
            style: Optional style name to apply
        """
        if not self.document:
            raise EditOperationError("No document loaded")
        
        paragraph = self.document.add_paragraph(text)
        if style and style in self.original_styles:
            paragraph.style = self.original_styles[style]['style']
    
    def get_document_text(self) -> str:
        """
        Get all text content from the document
        
        Returns:
            Complete text content of the document
        """
        if not self.document:
            return ""
        
        full_text = []
        for paragraph in self.document.paragraphs:
            full_text.append(paragraph.text)
        
        return '\n'.join(full_text)
    
    def get_document_info(self) -> Dict[str, Any]:
        """
        Get information about the document
        
        Returns:
            Dictionary containing document metadata
        """
        if not self.document:
            return {}
        
        return {
            'paragraph_count': len(self.document.paragraphs),
            'table_count': len(self.document.tables),
            'style_count': len(self.original_styles),
            'styles': list(self.original_styles.keys())
        }
    
    def save_to_buffer(self) -> io.BytesIO:
        """
        Save the document to a bytes buffer
        
        Returns:
            BytesIO buffer containing the modified DOCX document
            
        Raises:
            EditOperationError: If saving fails
        """
        if not self.document:
            raise EditOperationError("No document loaded")
        
        try:
            buffer = io.BytesIO()
            self.document.save(buffer)
            buffer.seek(0)
            return buffer
        except Exception as e:
            raise EditOperationError(f"Failed to save document: {str(e)}")


def create_sample_docx() -> io.BytesIO:
    """
    Create a sample DOCX document for testing purposes
    
    Returns:
        BytesIO buffer containing a sample DOCX document
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample Document', level=1)
    
    # Add paragraphs with different formatting
    p1 = doc.add_paragraph('This is a ')
    run1 = p1.add_run('bold')
    run1.bold = True
    p1.add_run(' word in a sentence.')
    
    p2 = doc.add_paragraph('This paragraph has some ')
    run2 = p2.add_run('italic text')
    run2.italic = True
    p2.add_run(' and some ')
    run3 = p2.add_run('underlined text')
    run3.underline = True
    p2.add_run('.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Name'
    table.cell(0, 1).text = 'Value'
    table.cell(1, 0).text = 'Sample Item'
    table.cell(1, 1).text = 'Sample Value'
    
    # Add another paragraph
    doc.add_paragraph('This is a paragraph that can be modified during testing.')
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer