"""
Resume-specific DOCX processing module for dynamic content replacement
"""

import io
import json
from typing import Dict, List, Any, Optional, Union
from datetime import datetime

from .docx_processor import DocxProcessor
from .exceptions import EditOperationError


class ResumeProcessor(DocxProcessor):
    """
    Specialized processor for resume documents with JSON data replacement
    """
    
    def __init__(self):
        super().__init__()
        self.resume_data: Optional[Dict[str, Any]] = None
    
    def set_resume_data(self, resume_json: Union[str, Dict[str, Any]]) -> None:
        """
        Set the resume data from JSON string or dictionary
        
        Args:
            resume_json: JSON string or dictionary containing resume data
        """
        if isinstance(resume_json, str):
            try:
                self.resume_data = json.loads(resume_json)
            except json.JSONDecodeError as e:
                raise EditOperationError(f"Invalid JSON format: {str(e)}")
        else:
            self.resume_data = resume_json
    
    def apply_resume_data(self, field_mappings: Optional[Dict[str, str]] = None) -> int:
        """
        Apply resume data to the document using field mappings
        
        Args:
            field_mappings: Optional custom field mappings for placeholder replacement
            
        Returns:
            Number of replacements made
        """
        if not self.resume_data:
            raise EditOperationError("No resume data provided")
        
        # Default field mappings (placeholder -> JSON path)
        default_mappings = {
            "{{NAME}}": "name",
            "{{PHONE}}": "contact.phone", 
            "{{EMAIL}}": "contact.email",
            "{{LINKEDIN}}": "contact.linkedin",
            "{{PROFESSIONAL_SUMMARY}}": "professional_summary",
            "{{CURRENT_DATE}}": None,  # Special case for current date
        }
        
        # Use provided mappings or defaults
        mappings = field_mappings or default_mappings
        
        total_replacements = 0
        
        for placeholder, json_path in mappings.items():
            try:
                if json_path is None:
                    # Handle special cases
                    if placeholder == "{{CURRENT_DATE}}":
                        value = datetime.now().strftime("%B %Y")
                    else:
                        continue
                else:
                    # Get value from nested JSON path
                    value = self._get_nested_value(self.resume_data, json_path)
                
                if value is not None:
                    replacements = self.find_and_replace_text(
                        search_text=placeholder,
                        replace_text=str(value),
                        preserve_formatting=True,
                        case_sensitive=True
                    )
                    total_replacements += replacements
                    
            except Exception as e:
                # Log error but continue with other replacements
                print(f"Warning: Failed to replace {placeholder}: {str(e)}")
        
        return total_replacements
    
    def apply_technical_skills(self, placeholder: str = "{{TECHNICAL_SKILLS}}") -> int:
        """
        Replace technical skills placeholder with formatted skills list
        
        Args:
            placeholder: The placeholder text to replace
            
        Returns:
            Number of replacements made
        """
        if not self.resume_data or "technical_skills" not in self.resume_data:
            return 0
        
        skills = self.resume_data["technical_skills"]
        if not isinstance(skills, list):
            return 0
        
        # Format skills as comma-separated list
        skills_text = ", ".join(skills)
        
        return self.find_and_replace_text(
            search_text=placeholder,
            replace_text=skills_text,
            preserve_formatting=True,
            case_sensitive=True
        )
    
    def apply_professional_experience(self, placeholder: str = "{{PROFESSIONAL_EXPERIENCE}}") -> int:
        """
        Replace professional experience placeholder with formatted experience
        
        Args:
            placeholder: The placeholder text to replace
            
        Returns:
            Number of replacements made
        """
        if not self.resume_data or "professional_experience" not in self.resume_data:
            return 0
        
        experiences = self.resume_data["professional_experience"]
        if not isinstance(experiences, list):
            return 0
        
        # Build formatted experience content
        experience_content = []
        
        for exp in experiences:
            # Add company and title line
            company_title = f"{exp.get('company', '')} - {exp.get('title', '')}"
            experience_content.append(company_title)
            
            # Add location and duration line  
            location_duration = f"{exp.get('location', '')} | {exp.get('duration', '')}"
            experience_content.append(location_duration)
            
            # Add highlights as bullet points
            highlights = exp.get('highlights', [])
            if highlights:
                for highlight in highlights:
                    experience_content.append(f"• {highlight}")
            
            # Add empty line between experiences
            experience_content.append("")
        
        # Join all content and replace
        experience_text = "\n".join(experience_content).strip()
        
        # First, replace the placeholder with basic text
        replacements = self.find_and_replace_text(
            search_text=placeholder,
            replace_text=experience_text,
            preserve_formatting=True,
            case_sensitive=True
        )
        
        # Then apply bullet formatting to lines that start with "•"
        if replacements > 0:
            self._apply_bullet_formatting_to_document()
        
        return replacements
    
    def _apply_bullet_formatting_to_document(self) -> None:
        """Apply bullet formatting to all paragraphs that contain bullet characters"""
        if not self.document:
            return
            
        # Find paragraphs that contain bullets and split them properly
        paragraphs_to_process = []
        standalone_bullets = []
        
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text.strip()
            if '•' in text and '\n' in text:
                # Multi-line bullets that need splitting
                paragraphs_to_process.append((i, paragraph))
            elif text.startswith('•'):
                # Single-line bullets that need formatting
                standalone_bullets.append(paragraph)
        
        # Process multi-line bullets in reverse order to avoid index shifting
        for para_index, paragraph in reversed(paragraphs_to_process):
            self._split_paragraph_with_bullets(paragraph, para_index)
            
        # Apply formatting to standalone bullets
        for paragraph in standalone_bullets:
            self._apply_bullet_formatting(paragraph)
            
    def _split_paragraph_with_bullets(self, paragraph, para_index: int) -> None:
        """Split a paragraph containing bullets into separate paragraphs with proper formatting"""
        if not self.document:
            return
            
        text = paragraph.text
        if '•' not in text or '\n' not in text:
            return
            
        # Split the text into lines
        lines = text.split('\n')
        if len(lines) <= 1:
            return
            
        # Clear the original paragraph and add the first line
        paragraph.clear()
        first_line = lines[0].strip()
        if first_line:
            paragraph.add_run(first_line)
        
        # Get the parent element for inserting new paragraphs
        parent_element = paragraph._element.getparent()
        paragraph_element_index = list(parent_element).index(paragraph._element)
        
        # Add subsequent lines as new paragraphs
        for i, line in enumerate(lines[1:], 1):
            line = line.strip()
            if not line:  # Skip empty lines
                continue
                
            # Create new paragraph
            new_para = self.document.add_paragraph()
            new_para.add_run(line)
            
            # Apply bullet formatting if this line starts with bullet
            if line.startswith('•'):
                self._apply_bullet_formatting(new_para)
            
            # Insert the new paragraph after the original
            parent_element.insert(paragraph_element_index + i, new_para._element)
    
    def apply_projects(self, placeholder: str = "{{PROJECTS}}") -> int:
        """
        Replace projects placeholder with formatted projects list
        
        Args:
            placeholder: The placeholder text to replace
            
        Returns:
            Number of replacements made
        """
        if not self.resume_data or "projects" not in self.resume_data:
            return 0
        
        projects = self.resume_data["projects"]
        if not isinstance(projects, list):
            return 0
        
        # Build formatted projects content
        projects_content = []
        
        for project in projects:
            # Add project name
            project_name = project.get('name', '')
            projects_content.append(project_name)
            
            # Add project highlights as bullet points
            highlights = project.get('highlights', [])
            if highlights:
                for highlight in highlights:
                    projects_content.append(f"• {highlight}")
            
            # Add empty line between projects
            projects_content.append("")
        
        # Join all content and replace
        projects_text = "\n".join(projects_content).strip()
        
        # First, replace the placeholder with basic text
        replacements = self.find_and_replace_text(
            search_text=placeholder,
            replace_text=projects_text,
            preserve_formatting=True,
            case_sensitive=True
        )
        
        # Then apply bullet formatting to lines that start with "•"
        if replacements > 0:
            self._apply_bullet_formatting_to_document()
        
        return replacements
    
    def apply_education(self, placeholder: str = "{{EDUCATION}}") -> int:
        """
        Replace education placeholder with formatted education list
        
        Args:
            placeholder: The placeholder text to replace
            
        Returns:
            Number of replacements made
        """
        if not self.resume_data or "education" not in self.resume_data:
            return 0
        
        education = self.resume_data["education"]
        if not isinstance(education, list):
            return 0
        
        # Format education entries
        education_text = ""
        for edu in education:
            education_text += f"{edu.get('degree', '')} - {edu.get('institution', '')}\n"
            education_text += f"{edu.get('duration', '')}\n\n"
        
        return self.find_and_replace_text(
            search_text=placeholder,
            replace_text=education_text.strip(),
            preserve_formatting=True,
            case_sensitive=True
        )
    
    def apply_content_replacement(self, content_mappings: Optional[Dict[str, str]] = None) -> Dict[str, int]:
        """
        Replace actual content in the document with new values from resume data
        
        Args:
            content_mappings: Optional mapping of current content to JSON paths
                             e.g., {"Ganesh Sarakadam": "name", "ganesh@email.com": "contact.email"}
            
        Returns:
            Dictionary showing number of replacements made for each category
        """
        if not self.resume_data:
            raise EditOperationError("No resume data provided")
        
        replacements_made = {
            'name_replacements': 0,
            'contact_replacements': 0,
            'content_replacements': 0,
            'total_replacements': 0
        }
        
        # If custom mappings provided, use them
        if content_mappings:
            for old_text, json_path in content_mappings.items():
                new_value = self._get_nested_value(self.resume_data, json_path)
                if new_value:
                    count = self.find_and_replace_text(
                        search_text=old_text,
                        replace_text=str(new_value),
                        preserve_formatting=True,
                        case_sensitive=False  # More flexible matching
                    )
                    replacements_made['content_replacements'] += count
                    replacements_made['total_replacements'] += count
        else:
            # Auto-detect and replace common resume fields
            replacements_made.update(self._auto_replace_resume_content())
        
        return replacements_made
    
    def _auto_replace_resume_content(self) -> Dict[str, int]:
        """
        Automatically detect and replace common resume content
        """
        replacements = {
            'name_replacements': 0,
            'contact_replacements': 0,
            'content_replacements': 0,
            'total_replacements': 0
        }
        
        # Get document text to analyze
        doc_text = self.get_document_text()
        
        # Replace name if provided
        if 'name' in self.resume_data:
            new_name = self.resume_data['name']
            
            # Try to find existing names in the document
            # Look for common name patterns (capitalized words)
            import re
            name_patterns = re.findall(r'\b[A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', doc_text)
            
            for potential_name in name_patterns:
                # Skip common non-name patterns
                if any(word in potential_name.lower() for word in ['company', 'corp', 'inc', 'llc', 'university', 'college']):
                    continue
                
                count = self.find_and_replace_text(
                    search_text=potential_name,
                    replace_text=new_name,
                    preserve_formatting=True,
                    case_sensitive=False
                )
                if count > 0:
                    replacements['name_replacements'] += count
                    replacements['total_replacements'] += count
                    print(f"DEBUG: Replaced '{potential_name}' with '{new_name}' ({count} times)")
        
        # Replace contact information
        if 'contact' in self.resume_data:
            contact = self.resume_data['contact']
            
            # Replace email
            if 'email' in contact:
                new_email = contact['email']
                # Find email patterns
                email_patterns = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', doc_text)
                for old_email in email_patterns:
                    count = self.find_and_replace_text(
                        search_text=old_email,
                        replace_text=new_email,
                        preserve_formatting=True,
                        case_sensitive=False
                    )
                    if count > 0:
                        replacements['contact_replacements'] += count
                        replacements['total_replacements'] += count
                        print(f"DEBUG: Replaced email '{old_email}' with '{new_email}' ({count} times)")
            
            # Replace phone
            if 'phone' in contact:
                new_phone = contact['phone']
                # Find phone patterns
                phone_patterns = re.findall(r'[\+]?[1-9]?[\d\s\-\(\)\.]{10,}', doc_text)
                for old_phone in phone_patterns:
                    # Skip if it looks like a year or other number
                    if len(old_phone.replace(' ', '').replace('-', '').replace('(', '').replace(')', '').replace('.', '')) >= 10:
                        count = self.find_and_replace_text(
                            search_text=old_phone.strip(),
                            replace_text=new_phone,
                            preserve_formatting=True,
                            case_sensitive=False
                        )
                        if count > 0:
                            replacements['contact_replacements'] += count
                            replacements['total_replacements'] += count
                            print(f"DEBUG: Replaced phone '{old_phone}' with '{new_phone}' ({count} times)")
        
        return replacements

    def apply_json_template_replacement(self, custom_mappings: Optional[Dict[str, str]] = None) -> Dict[str, int]:
        """
        Replace placeholders in the document with values from JSON data
        Supports nested JSON paths like {{contact.email}}, {{professional_experience.0.company}}
        Also supports custom mappings for cleaner template placeholders
        
        Args:
            custom_mappings: Optional mapping of short placeholders to JSON paths
                           e.g., {"company1": "professional_experience.0.company"}
        
        Returns:
            Dictionary showing number of replacements made for each category
        """
        if not self.resume_data:
            raise EditOperationError("No resume data provided")
        
        replacements_made = {
            'basic_fields': 0,
            'contact_fields': 0,
            'technical_skills': 0,
            'professional_experience': 0,
            'projects': 0,
            'education': 0,
            'total_replacements': 0
        }
        
        # Get document text to find all placeholders
        doc_text = self.get_document_text()
        
        # Find all placeholder patterns {{...}}
        import re
        placeholder_pattern = r'\{\{([^}]+)\}\}'
        placeholders = re.findall(placeholder_pattern, doc_text)
        
        print(f"DEBUG: Found placeholders in document: {placeholders}")
        
        # Create default short mappings for common long placeholders
        default_mappings = self._get_default_short_mappings()
        
        # Merge custom mappings with defaults
        all_mappings = {**default_mappings}
        if custom_mappings:
            all_mappings.update(custom_mappings)
        
        for placeholder in placeholders:
            placeholder_key = placeholder.strip()
            
            # Check if it's a short placeholder that needs mapping
            if placeholder_key in all_mappings:
                json_path = all_mappings[placeholder_key]
                replacement_value = self._get_nested_json_value(self.resume_data, json_path)
                print(f"DEBUG: Short mapping '{placeholder_key}' -> '{json_path}' = '{replacement_value}'")
            else:
                # Direct JSON path
                replacement_value = self._get_nested_json_value(self.resume_data, placeholder_key)
            
            if replacement_value is not None:
                # Check if this is a highlights/bullet point replacement
                full_placeholder = "{{" + placeholder + "}}"
                
                if (isinstance(replacement_value, list) and 
                    any(term in placeholder_key.lower() for term in ['highlights', 'responsibilities'])):
                    # Use specialized bullet point replacement for better spacing
                    count = self.replace_with_bullet_points(
                        search_text=full_placeholder,
                        bullet_items=[str(item) for item in replacement_value],
                        preserve_formatting=True
                    )
                else:
                    # Convert value to string format for regular replacement
                    formatted_value = self._format_json_value(replacement_value, placeholder_key)
                    
                    # Perform regular replacement
                    count = self.find_and_replace_text(
                        search_text=full_placeholder,
                        replace_text=formatted_value,
                        preserve_formatting=True,
                        case_sensitive=False
                    )
                
                if count > 0:
                    # Categorize the replacement
                    category = self._categorize_placeholder(placeholder_key)
                    replacements_made[category] += count
                    replacements_made['total_replacements'] += count
                    
                    # Create appropriate debug message
                    if (isinstance(replacement_value, list) and 
                        any(term in placeholder_key.lower() for term in ['highlights', 'responsibilities'])):
                        print(f"DEBUG: Replaced '{full_placeholder}' with {len(replacement_value)} bullet points ({count} times)")
                    else:
                        formatted_value = self._format_json_value(replacement_value, placeholder_key)
                        print(f"DEBUG: Replaced '{full_placeholder}' with '{formatted_value}' ({count} times)")
                else:
                    print(f"DEBUG: Placeholder '{full_placeholder}' not found in document")
            else:
                print(f"DEBUG: No value found for placeholder '{placeholder_key}' in JSON data")
        
        # No need to remove placeholder bullet points since they're not in the template anymore
        
        # Apply bullet formatting to any paragraphs containing bullets
        self._apply_bullet_formatting_to_document()
        
        return replacements_made
    
    def _get_default_short_mappings(self) -> Dict[str, str]:
        """
        Get default short placeholder mappings for cleaner templates
        """
        return {
            # Basic fields
            "{{NAME}}": "name",  # Support uppercase NAME
            "{{name}}": "name",
            "{{professional_summary}}": "professional_summary",
            "{{PROFESSIONAL_SUMMARY}}": "professional_summary",  # Uppercase version
            "{{TECHNICAL_SKILLS}}": "technical_skills",  # Uppercase version
            "{{technical_skills}}": "technical_skills",
            # Note: PROFESSIONAL_EXPERIENCE and PROJECTS are handled by specialized methods
            # Note: EDUCATION is handled by specialized method
            
            # Experience mappings (simplified - only highlights)
            "{{highlights1}}": "highlights1",
            "{{highlights2}}": "highlights2", 
            "{{highlights3}}": "highlights3",
            
            # Project mappings (simplified - name and highlights only)
            "{{projects1}}": "projects1",
            "{{projects2}}": "projects2",
            "{{project1_highlights}}": "project1_highlights",
            "{{project2_highlights}}": "project2_highlights",
            
            # Education mappings (if needed)
            "{{education}}": "education",
        }
    
    def _get_nested_json_value(self, data: Dict[str, Any], path: str) -> Any:
        """
        Get value from nested JSON using dot notation or array indices
        Examples: 'name', 'contact.email', 'professional_experience.0.company'
        """
        try:
            keys = path.split('.')
            value = data
            
            for key in keys:
                if key.isdigit():  # Array index
                    if isinstance(value, list) and int(key) < len(value):
                        value = value[int(key)]
                    else:
                        return None
                else:  # Object key
                    if isinstance(value, dict) and key in value:
                        value = value[key]
                    else:
                        return None
            
            return value
        except (KeyError, IndexError, TypeError, ValueError):
            return None
    
    def _format_json_value(self, value: Any, placeholder_key: str) -> str:
        """
        Format JSON values for document insertion
        """
        if isinstance(value, list):
            if 'technical_skills' in placeholder_key.lower():
                # Format technical skills as comma-separated
                return ', '.join(str(item) for item in value)
            elif any(term in placeholder_key.lower() for term in ['highlights', 'responsibilities']):
                # Format highlights/responsibilities as bullet points with controlled spacing
                if not value:
                    return ""
                
                # Create bullet points with proper spacing for Word documents
                # Each bullet point should be on its own line with consistent formatting
                bullet_points = []
                for item in value:
                    bullet_points.append(f"• {str(item)}")
                
                # Join with line breaks for proper spacing in Word
                return '\n'.join(bullet_points)
            else:
                # Default list formatting
                return ', '.join(str(item) for item in value)
        elif isinstance(value, dict):
            # For dict values, you might want custom formatting
            return str(value)
        else:
            return str(value)
    
    def _categorize_placeholder(self, placeholder_key: str) -> str:
        """
        Categorize placeholder for tracking replacements
        """
        key_lower = placeholder_key.lower()
        
        if any(term in key_lower for term in ['contact', 'phone', 'email', 'linkedin']):
            return 'contact_fields'
        elif 'technical_skills' in key_lower:
            return 'technical_skills'
        elif 'professional_experience' in key_lower:
            return 'professional_experience'
        elif 'project' in key_lower:
            return 'projects'
        elif 'education' in key_lower:
            return 'education'
        else:
            return 'basic_fields'

    def apply_all_resume_data(self, custom_mappings: Optional[Dict[str, str]] = None) -> Dict[str, int]:
        """
        Apply all resume data sections to the document using placeholder replacement
        
        Args:
            custom_mappings: Optional custom field mappings
            
        Returns:
            Dictionary with section names and replacement counts
        """
        results = {}
        
        # Apply basic fields (placeholder-based)
        results["basic_fields"] = self.apply_resume_data(custom_mappings)
        
        # Apply technical skills
        results["technical_skills"] = self.apply_technical_skills()
        
        # Apply professional experience
        results["professional_experience"] = self.apply_professional_experience()
        
        # Apply projects
        results["projects"] = self.apply_projects()
        
        # Apply education
        results["education"] = self.apply_education()
        
        return results
    
    def _get_nested_value(self, data: Dict[str, Any], path: str) -> Any:
        """
        Get value from nested dictionary using dot notation
        
        Args:
            data: Dictionary to search in
            path: Dot-separated path (e.g., "contact.email")
            
        Returns:
            Value at the specified path or None if not found
        """
        try:
            keys = path.split('.')
            value = data
            
            for key in keys:
                if isinstance(value, dict) and key in value:
                    value = value[key]
                else:
                    return None
            
            return value
        except (KeyError, TypeError, AttributeError):
            return None
    
    def get_resume_placeholders(self) -> List[str]:
        """
        Get list of common resume placeholders that can be used in templates
        
        Returns:
            List of placeholder strings
        """
        return [
            "{{NAME}}",
            "{{PHONE}}",
            "{{EMAIL}}",
            "{{LINKEDIN}}",
            "{{PROFESSIONAL_SUMMARY}}",
            "{{TECHNICAL_SKILLS}}",
            "{{PROFESSIONAL_EXPERIENCE}}",
            "{{PROJECTS}}",
            "{{EDUCATION}}",
            "{{CURRENT_DATE}}"
        ]
    
    def validate_resume_data(self) -> Dict[str, Any]:
        """
        Validate the resume data structure
        
        Returns:
            Dictionary with validation results
        """
        if not self.resume_data:
            return {"valid": False, "errors": ["No resume data provided"]}
        
        errors = []
        warnings = []
        
        # Required fields
        required_fields = ["name", "contact", "professional_summary"]
        for field in required_fields:
            if field not in self.resume_data:
                errors.append(f"Missing required field: {field}")
        
        # Validate contact information
        if "contact" in self.resume_data:
            contact = self.resume_data["contact"]
            contact_fields = ["phone", "email"]
            for field in contact_fields:
                if field not in contact:
                    warnings.append(f"Missing contact field: {field}")
        
        # Validate arrays
        array_fields = ["technical_skills", "professional_experience", "projects", "education"]
        for field in array_fields:
            if field in self.resume_data and not isinstance(self.resume_data[field], list):
                errors.append(f"Field {field} should be an array")
        
        return {
            "valid": len(errors) == 0,
            "errors": errors,
            "warnings": warnings
        }


def create_resume_template() -> io.BytesIO:
    """
    Create a sample resume template with placeholders
    
    Returns:
        BytesIO buffer containing a resume template DOCX document
    """
    from docx import Document
    from docx.shared import Inches
    
    doc = Document()
    
    # Header with name and contact
    header = doc.add_heading('{{NAME}}', level=1)
    
    contact_para = doc.add_paragraph()
    contact_para.add_run('Phone: {{PHONE}} | Email: {{EMAIL}} | LinkedIn: {{LINKEDIN}}')
    
    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=2)
    doc.add_paragraph('{{PROFESSIONAL_SUMMARY}}')
    
    # Technical Skills
    doc.add_heading('TECHNICAL SKILLS', level=2)
    doc.add_paragraph('{{TECHNICAL_SKILLS}}')
    
    # Professional Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
    doc.add_paragraph('{{PROFESSIONAL_EXPERIENCE}}')
    
    # Projects
    doc.add_heading('PROJECTS', level=2)
    doc.add_paragraph('{{PROJECTS}}')
    
    # Education
    doc.add_heading('EDUCATION', level=2)
    doc.add_paragraph('{{EDUCATION}}')
    
    # Footer
    doc.add_paragraph(f'\nGenerated on {{CURRENT_DATE}}')
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer